"""Production PDF, DOCX and XLSX renderers."""
from pathlib import Path
from typing import Any
ENGINE_VERSION="3.7.0"

def render_pdf(model:dict[str,Any],target:str|Path)->str:
 from reportlab.lib.pagesizes import A4
 from reportlab.platypus import SimpleDocTemplate,Paragraph,Spacer,Table,PageBreak
 from reportlab.lib.styles import getSampleStyleSheet
 target=Path(target); target.parent.mkdir(parents=True,exist_ok=True); styles=getSampleStyleSheet(); story=[]
 for s in model.get("sections",[]):
  c=s.get("content",""); k=s.get("type")
  if k=="cover": story+=[Paragraph("VvE Navigator",styles["Title"]),Paragraph(str(c.get("vve_name","")),styles["Heading1"]),PageBreak()]; continue
  story.append(Paragraph(str(s.get("title","")),styles["Heading2"]))
  if k=="table": story.append(Table([[str(r.get("kpi","")),str(r.get("value",""))] for r in c] or [["",""]]))
  elif k=="list": [story.append(Paragraph(f"• {x}",styles["Normal"])) for x in c]
  elif isinstance(c,dict): [story.append(Paragraph(str(v),styles["Normal"])) for v in c.values()]
  else: story.append(Paragraph(str(c),styles["Normal"]))
  story.append(Spacer(1,10))
 SimpleDocTemplate(str(target),pagesize=A4,title=model.get("document_title","VvE Navigator")).build(story); return str(target)

def render_docx(model:dict[str,Any],target:str|Path)->str:
 from docx import Document
 target=Path(target); target.parent.mkdir(parents=True,exist_ok=True); doc=Document(); doc.core_properties.title=model.get("document_title","VvE Navigator")
 for s in model.get("sections",[]):
  c=s.get("content",""); k=s.get("type"); doc.add_heading(str(s.get("title","")),0 if k=="cover" else 1)
  if k=="table":
   t=doc.add_table(rows=1,cols=2); t.rows[0].cells[0].text="KPI"; t.rows[0].cells[1].text="Waarde"
   for r in c: cells=t.add_row().cells; cells[0].text=str(r.get("kpi","")); cells[1].text=str(r.get("value",""))
  elif k=="list": [doc.add_paragraph(str(x),style="List Bullet") for x in c]
  elif isinstance(c,dict): [doc.add_paragraph(str(v)) for v in c.values()]
  else: doc.add_paragraph(str(c))
 doc.save(str(target)); return str(target)

def render_xlsx(model:dict[str,Any],target:str|Path)->str:
 from openpyxl import Workbook
 target=Path(target); target.parent.mkdir(parents=True,exist_ok=True); wb=Workbook(); wb.remove(wb.active)
 for s in model.get("sections",[]):
  ws=wb.create_sheet((str(s.get("title","Sheet")) or "Sheet")[:31]); c=s.get("content",""); k=s.get("type"); ws["A1"]=s.get("title","")
  if k=="table":
   ws.append(["KPI","Waarde"]); [ws.append([r.get("kpi",""),r.get("value","")]) for r in c]
  elif k=="list": [ws.append([x]) for x in c]
  elif isinstance(c,dict): [ws.append([a,str(b)]) for a,b in c.items()]
  else: ws["A3"]=str(c)
 wb.save(str(target)); return str(target)

def render_all(model:dict[str,Any],files:dict[str,str])->dict[str,str]:
 return {"pdf":render_pdf(model,files["pdf"]),"docx":render_docx(model,files["docx"]),"xlsx":render_xlsx(model,files["xlsx"])}
